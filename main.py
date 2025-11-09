import tkinter as tk
from tkinter import filedialog, messagebox
from docx import Document


def choose_file():
    global chosen_file
    file_path = filedialog.askopenfilename(filetypes=[("Wordownik", "*.docx")])
    if file_path:
        chosen_file = file_path
        label_info.config(text=f"File:\n{file_path.split('/')[-1]}")
        frame_function.pack(pady=20)  # show function buttons


def change_word():
    if not chosen_file:
        messagebox.showerror("Error", "First pick a file.")
        return

    doc = Document(chosen_file)
    old_word = entry_old.get()
    new_word = entry_new.get()

    if not old_word or not new_word:
        messagebox.showerror("Error", "Insert a new word and choose word to pe replaced.")
        return

    for p in doc.paragraphs:
        full_text = "".join(run.text for run in p.runs)
        if old_word in full_text:
            new_text = full_text.replace(old_word, new_word)
            for run in p.runs:
                run.text = ""
            p.runs[0].text = new_text

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    full_text = "".join(run.text for run in p.runs)
                    if old_word in full_text:
                        new_text = full_text.replace(old_word, new_word)
                        for run in p.runs:
                            run.text = ""
                        p.runs[0].text = new_text

    save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
    if save_path:
        doc.save(save_path)
        messagebox.showinfo("Action", "The replacement was made and the file was saved!")


def counting():
    if not chosen_file:
        messagebox.showerror("Error", "First pick a file.")
        return

    doc = Document(chosen_file)
    text = " ".join(p.text for p in doc.paragraphs)
    number = len(text.split())
    messagebox.showinfo("Counting", f"Number of words in the document: {number}")


# --- GUI ---
root = tk.Tk()
root.title("Wordownik")
root.configure(bg="#ACC7B4")

chosen_file = None

# Choose file
frame_choice = tk.Frame(root, bg="#ACC7B4")
frame_choice.pack(pady=20)

btn_choice = tk.Button(frame_choice, text="Choose Word file", command=choose_file,
                        bg="#D0E6B9", fg="#331B3F", font=("Montserrat", 10, "bold"), width=25)
btn_choice.pack()

label_info = tk.Label(frame_choice, text="No file selected", bg="#ACC7B4", fg="#331B3F", font=("Montserrat", 9))
label_info.pack(pady=10)

# function (hidden)
frame_function = tk.Frame(root, bg="#ACC7B4")

# --- Replace ---
tk.Label(frame_function, text="Word to replace:", bg="#ACC7B4", fg="#331B3F", font=("Montserrat", 10, "bold")).pack(pady=5)
entry_old = tk.Entry(frame_function, width=30)
entry_old.pack(pady=5)

tk.Label(frame_function, text="New word:", bg="#ACC7B4", fg="#331B3F", font=("Montserrat", 10, "bold")).pack(pady=5)
entry_new = tk.Entry(frame_function, width=30)
entry_new.pack(pady=5)

btn_replace = tk.Button(frame_function, text="Change a word", command=change_word,
                       bg="#D0E6B9", fg="#331B3F", font=("Montserrat", 10, "bold"), width=20)
btn_replace.pack(pady=10)

# --- Inna funkcja ---
btn_counting = tk.Button(frame_function, text="Count the words", command=counting,
                       bg="#D0E6B9", fg="#331B3F", font=("Montserrat", 10, "bold"), width=20)
btn_counting.pack(pady=10)

root.mainloop()